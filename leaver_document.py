from __future__ import annotations

from datetime import date
from pathlib import Path
import re

from docx import Document


_TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "leaver_return_template.docx"
_GBS_TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "leaver_gbs_template.docx"
_POZNAN_TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "leaver_poznan_template.docx"
_OUTPUT_DIR = Path(__file__).resolve().parent / "generated_docs"
_RECEIVER_ROLE = "IT engineer"
_GBS_RECEIVER_ROLE = "Senior IT Engineer Service Desk"
_POZNAN_RECEIVER_ROLE = "IT engineer"

_ASSET_CATEGORY_ORDER = {
    "laptop": 0,
    "notebook": 0,
    "monitor": 1,
    "phone": 2,
    "mobile": 2,
    "sim": 3,
    "headset": 4,
    "keyboard": 5,
    "mouse": 6,
    "dock": 7,
}


def _format_date(value) -> str:
    if isinstance(value, date):
        return value.isoformat()
    return str(value or "").strip()


def _format_date_dmy(value) -> str:
    """Return DD.MM.YYYY format used in GBS documents."""
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    s = str(value or "").strip()
    if re.match(r'^\d{4}-\d{2}-\d{2}$', s):
        try:
            return date.fromisoformat(s).strftime("%d.%m.%Y")
        except Exception:
            pass
    return s


def _receiver_initials(name: str) -> str:
    parts = name.strip().split()
    if len(parts) >= 2:
        return f"{parts[0][0]}. {' '.join(parts[1:])}"
    return name


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\\\|?*]+', "_", value or "")
    cleaned = re.sub(r"\s+", " ", cleaned).strip().rstrip(".")
    return cleaned or "leaver"


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def _asset_name(asset: dict) -> str:
    return (
        (asset.get("category") or {}).get("name")
        or asset.get("name")
        or (asset.get("model") or {}).get("name")
        or "IT asset"
    )


def _manufacturer_model(asset: dict) -> str:
    model = (asset.get("model") or {}).get("name", "").strip()
    manufacturer = ((asset.get("model") or {}).get("manufacturer") or {}).get("name", "").strip()
    if manufacturer and model:
        if manufacturer.casefold() in model.casefold():
            return model
        return f"{manufacturer} {model}"
    return model or manufacturer or asset.get("name", "")


def _asset_sort_key(asset: dict) -> tuple[int, str, str]:
    category = ((asset.get("category") or {}).get("name") or "").casefold()
    order = 99
    for keyword, rank in _ASSET_CATEGORY_ORDER.items():
        if keyword in category:
            order = rank
            break
    return (order, category, (asset.get("asset_tag") or asset.get("name") or "").casefold())


def _ensure_asset_rows(table, wanted_rows: int) -> None:
    while len(table.rows) - 1 < wanted_rows:
        table.add_row()


def _fill_asset_table(table, assets: list[dict]) -> None:
    _ensure_asset_rows(table, max(4, len(assets)))

    row_count = len(table.rows) - 1
    for idx in range(row_count):
        cells = table.rows[idx + 1].cells
        if idx < len(assets):
            asset = assets[idx]
            cells[0].text = str(idx + 1)
            cells[1].text = _asset_name(asset)
            cells[2].text = _manufacturer_model(asset)
            cells[3].text = (asset.get("asset_tag") or asset.get("serial") or "").strip()
            cells[4].text = "+"
            cells[5].text = "Def. nėra"
        else:
            for cell in cells:
                cell.text = ""


def _fill_gbs_date_table(table, document_date: str) -> None:
    """Fill the Date/Place header table in the GBS template."""
    table.rows[1].cells[1].text = document_date
    table.rows[2].cells[1].text = "GBS Hub"


def _fill_gbs_asset_table(table, assets: list[dict]) -> None:
    _ensure_asset_rows(table, max(4, len(assets)))
    row_count = len(table.rows) - 1
    for idx in range(row_count):
        cells = table.rows[idx + 1].cells
        if idx < len(assets):
            asset = assets[idx]
            cells[0].text = str(idx + 1)
            cells[1].text = _manufacturer_model(asset)
            cells[2].text = (asset.get("serial") or "").strip()
            cells[3].text = (asset.get("asset_tag") or "").strip()
            cells[4].text = ""
            cells[5].text = ""
        else:
            for cell in cells:
                cell.text = ""


def generate_poznan_leaver_return_document(ticket: dict, snipeit=None) -> tuple[Path, list[str]]:
    if not _POZNAN_TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Poznan template not found: {_POZNAN_TEMPLATE_PATH}\n"
            "Place the Poznan return act template at that path to enable document generation."
        )

    warnings: list[str] = []
    snipe_assets: list[dict] = []

    if snipeit:
        try:
            user = snipeit.find_user(ticket.get("first_name", ""), ticket.get("last_name", ""))
            if user:
                snipe_assets = snipeit.get_user_assets(user["id"])
                if not snipe_assets:
                    warnings.append("No assigned Snipe-IT assets were found for this user.")
            else:
                warnings.append("Could not find this leaver in Snipe-IT.")
        except Exception as exc:
            warnings.append(f"Snipe-IT lookup failed: {exc}")
    else:
        warnings.append("Snipe-IT is not configured, so the asset list could not be filled.")

    assets = sorted(snipe_assets, key=_asset_sort_key)

    doc = Document(_POZNAN_TEMPLATE_PATH)

    warnings.append(
        "Poznan template paragraph indices have not been configured yet — "
        "verify the generated document and update generate_poznan_leaver_return_document() "
        "to match your template layout."
    )

    _fill_asset_table(doc.tables[0], assets)

    _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{ticket.get('key', 'leaver')}_{_safe_filename(ticket.get('name', 'leaver'))}_return_act.docx"
    output_path = _OUTPUT_DIR / filename
    doc.save(output_path)
    return output_path, warnings


def generate_gbs_leaver_return_document(ticket: dict, snipeit=None) -> tuple[Path, list[str]]:
    if not _GBS_TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"GBS template not found: {_GBS_TEMPLATE_PATH}")

    warnings: list[str] = []
    today = date.today()
    snipe_user: dict = {}
    snipe_assets: list[dict] = []

    if snipeit:
        try:
            user = snipeit.find_user(ticket.get("first_name", ""), ticket.get("last_name", ""))
            if user:
                snipe_user = snipeit.get_user_details(user["id"])
                snipe_assets = snipeit.get_user_assets(user["id"])
                if not snipe_assets:
                    warnings.append("No assigned Snipe-IT assets were found for this user.")
            else:
                warnings.append("Could not find this leaver in Snipe-IT.")
        except Exception as exc:
            warnings.append(f"Snipe-IT lookup failed: {exc}")
    else:
        warnings.append("Snipe-IT is not configured, so the asset list could not be filled.")

    assets = sorted(snipe_assets, key=_asset_sort_key)
    employee_role = ticket.get("job_title") or snipe_user.get("jobtitle") or ""
    receiver_name = ticket.get("assignee_name") or ""
    document_date = _format_date_dmy(today)
    return_date = _format_date_dmy(ticket.get("last_day") or today)

    doc = Document(_GBS_TEMPLATE_PATH)

    _fill_gbs_date_table(doc.tables[0], document_date)

    _replace_paragraph_text(
        doc.paragraphs[5],
        f"Employee\nName, surname: {ticket.get('name', '')}\n"
        f"Position: {employee_role}\n"
        f"Return date: {return_date}\n"
        f"Recipient\nName, surname: {receiver_name}\n"
        f"Position: {_GBS_RECEIVER_ROLE}",
    )

    _fill_gbs_asset_table(doc.tables[1], assets)

    _replace_paragraph_text(
        doc.paragraphs[20],
        f"Employee {ticket.get('name', '')}\n"
        "Name, surname: _______________________\n"
        "Date: ________________________________\n"
        "Contact details for receiving documents: ________________________________",
    )
    _replace_paragraph_text(
        doc.paragraphs[21],
        f"Recipient \n"
        f"Signature: {_receiver_initials(receiver_name)}\n"
        f"Name, surname: {receiver_name} \n"
        f"Date {document_date}",
    )

    _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{ticket.get('key', 'leaver')}_{_safe_filename(ticket.get('name', 'leaver'))}_return_act.docx"
    output_path = _OUTPUT_DIR / filename
    doc.save(output_path)
    return output_path, warnings


def generate_leaver_return_document(ticket: dict, snipeit=None) -> tuple[Path, list[str]]:
    if not _TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {_TEMPLATE_PATH}")

    warnings: list[str] = []
    today = date.today()
    snipe_user: dict = {}
    snipe_assets: list[dict] = []

    if snipeit:
        try:
            user = snipeit.find_user(ticket.get("first_name", ""), ticket.get("last_name", ""))
            if user:
                snipe_user = snipeit.get_user_details(user["id"])
                snipe_assets = snipeit.get_user_assets(user["id"])
                if not snipe_assets:
                    warnings.append("No assigned Snipe-IT assets were found for this user.")
            else:
                warnings.append("Could not find this leaver in Snipe-IT.")
        except Exception as exc:
            warnings.append(f"Snipe-IT lookup failed: {exc}")
    else:
        warnings.append("Snipe-IT is not configured, so the asset list could not be filled.")

    assets = sorted(snipe_assets, key=_asset_sort_key)
    employee_role = ticket.get("job_title") or snipe_user.get("jobtitle") or ""
    location = (
        ticket.get("office")
        or ((snipe_user.get("location") or {}).get("name") if snipe_user else "")
        or ticket.get("company")
        or ""
    )
    receiver_name = ticket.get("assignee_name") or ""
    document_date = _format_date(today)
    return_date = _format_date(ticket.get("last_day") or today)

    doc = Document(_TEMPLATE_PATH)

    _replace_paragraph_text(
        doc.paragraphs[3],
        f"Data: {document_date}\nVieta: {location}",
    )
    _replace_paragraph_text(
        doc.paragraphs[5],
        "Darbuotojo vardas, pavardė: "
        f"{ticket.get('name', '')}\n"
        f"Pareigos: {employee_role}\n"
        f"Grąžinimo data: {return_date}",
    )
    _replace_paragraph_text(
        doc.paragraphs[6],
        "Priimančio asmens vardas, pavardė: "
        f"{receiver_name}\n"
        f"Pareigos: {_RECEIVER_ROLE}",
    )
    _replace_paragraph_text(
        doc.paragraphs[16],
        "Darbuotojas (grąžinantis įrangą):\n"
        "Parašas: ____________________________\n"
        f"Vardas, pavardė: {ticket.get('name', '')}\n"
        f"Data: {return_date}",
    )
    _replace_paragraph_text(
        doc.paragraphs[18],
        "Priėmė (atsakingas asmuo):\n"
        "Parašas: ____________________________\n"
        f"Vardas, pavardė: {receiver_name}\n"
        f"Data: {document_date}",
    )

    _delete_paragraph(doc.paragraphs[17])
    _delete_paragraph(doc.paragraphs[10])

    _fill_asset_table(doc.tables[0], assets)

    _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{ticket.get('key', 'leaver')}_{_safe_filename(ticket.get('name', 'leaver'))}_return_act.docx"
    output_path = _OUTPUT_DIR / filename
    doc.save(output_path)
    return output_path, warnings
