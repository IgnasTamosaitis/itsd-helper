"""Build a Confluence-importable Word document from the Jira Reminders KB."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = "0C66E4"
DARK = "172B4D"
GRAY = "5E6C84"
LIGHT_BLUE = "E9F2FF"
TABLE_HEADER = "DEEBFF"
CODE_BG = "F4F5F7"


def _shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _add_hyperlink(paragraph, text: str, url: str):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))"
)


def _add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(DARK)
        else:
            link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                _add_hyperlink(paragraph, link.group(1), link.group(2))
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = document.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name, size, color in (
        ("Title", 24, ACCENT),
        ("Heading 1", 16, DARK),
        ("Heading 2", 12, ACCENT),
        ("Heading 3", 10, DARK),
    ):
        style = document.styles[name]
        style.font.name = "Segoe UI"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "KB Code" not in document.styles:
        style = document.styles.add_style("KB Code", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(8.5)
        style.font.color.rgb = RGBColor.from_string(DARK)
        style.paragraph_format.left_indent = Inches(0.15)
        style.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    run = footer.add_run("Jira Reminders KB  |  Version 1.4.0  |  4 August 2026")
    run.font.name = "Segoe UI"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)


def _add_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    _shade(cell, LIGHT_BLUE)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    _add_inline(paragraph, text)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_table(document: Document, lines: list[str]) -> None:
    rows = [
        [cell.strip() for cell in line.strip().strip("|").split("|")]
        for line in lines
    ]
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        for column_index in range(width):
            cell = table.cell(row_index, column_index)
            value = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            _add_inline(paragraph, value)
            if row_index == 0:
                _shade(cell, TABLE_HEADER)
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build(source: Path, output: Path) -> None:
    document = Document()
    _configure_document(document)
    document.core_properties.title = "Jira Reminders — Installation and First-Time Setup"
    document.core_properties.subject = "IT Service Desk knowledge base"
    document.core_properties.author = "IT Service Desk"

    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph()
        _add_inline(paragraph, " ".join(line.strip() for line in paragraph_buffer))
        paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            paragraph = document.add_paragraph(style="KB Code")
            paragraph.add_run("\n".join(code_lines))
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), CODE_BG)
            paragraph._p.get_or_add_pPr().append(shading)
            index += 1
            continue

        if stripped.startswith("#"):
            flush_paragraph()
            match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
            if match:
                level = len(match.group(1))
                style = "Title" if level == 1 else f"Heading {min(level - 1, 3)}"
                paragraph = document.add_paragraph(style=style)
                _add_inline(paragraph, match.group(2))
                if level == 1:
                    paragraph.add_run().add_break(WD_BREAK.LINE)
                    subtitle = paragraph.add_run(
                        "End-user installation, setup, daily use, and troubleshooting"
                    )
                    subtitle.font.name = "Segoe UI"
                    subtitle.font.size = Pt(10)
                    subtitle.font.bold = False
                    subtitle.font.color.rgb = RGBColor.from_string(GRAY)
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            callout = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                callout.append(lines[index].strip()[1:].strip())
                index += 1
            _add_callout(document, " ".join(callout))
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            _add_table(document, table_lines)
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            flush_paragraph()
            paragraph = document.add_paragraph(
                style="List Bullet" if bullet else "List Number"
            )
            _add_inline(paragraph, (bullet or numbered).group(1))
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    docs_dir = Path(__file__).resolve().parent
    source = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else docs_dir / "Jira-Reminders-KB.md"
    output = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else docs_dir / "Jira-Reminders-KB.docx"
    build(source, output)
    print(output)


if __name__ == "__main__":
    main()
